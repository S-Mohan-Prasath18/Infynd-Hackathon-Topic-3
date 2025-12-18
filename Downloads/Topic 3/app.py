import os
import time
import json
# Check dependencies
try:
    from werkzeug.utils import secure_filename
    from flask import Flask, request, jsonify, send_file, render_template_string
    from flask_cors import CORS
    import pandas as pd
    from rapidfuzz import fuzz
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    print(f"\nCRITICAL ERROR: Missing dependency '{e.name}'.")
    print("Please run: pip install -r requirements.txt\n")
    # For user visibility if run in an environment that prevents installation
    print("If you are seeing this, the agent could not install dependencies automatically.")
    print("Please run the command manually in your terminal.")
    exit(1)

# Import our DQ modules
from dq import validation, corrections, dedupe, job_role, metrics, ai_cleaner

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 512 * 1024 * 1024  # 512MB limit for huge datasets

# Initialize AI
# API Key is now set in run.bat or environment
ai = ai_cleaner.AICleaner()

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.route('/')
def index():
    return send_file('index.html')

@app.route('/api/analyze', methods=['POST'])
def analyze():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file:
        timestamp = int(time.time())
        filename = secure_filename(file.filename)
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_{filename}")
        file.save(save_path)
        
        try:
            # Handle ZIP or CSV
            # If ZIP, extract the first CSV found
            if filename.lower().endswith('.zip'):
                import zipfile
                with zipfile.ZipFile(save_path, 'r') as zip_ref:
                    # Find first CSV
                    csv_files = [f for f in zip_ref.namelist() if f.lower().endswith('.csv')]
                    if not csv_files:
                        return jsonify({'error': 'No CSV file found in ZIP archive'}), 400
                    
                    # Extract to upload folder
                    target_csv = csv_files[0]
                    zip_ref.extract(target_csv, app.config['UPLOAD_FOLDER'])
                    save_path = os.path.join(app.config['UPLOAD_FOLDER'], target_csv)
                    print(f"[Core] Extracted {target_csv} from ZIP.")

            # Read CSV
            # optimize types for less memory if possible, but standard is fine for <500MB
            # Enforce low memory if needed, but for now standard read
            df = pd.read_csv(save_path)
            
            if df.empty:
                return jsonify({'error': 'CSV file is empty'}), 400
            
            # 1. Pre-Scan (AI Discovery for disguised metrics)
            ai_mapping = ai.identify_columns(df)
            
            # 2. Validation
            df = validation.validate_data(df, ai_mapping=ai_mapping)
            
            # 2. Corrections
            df = corrections.apply_corrections(df)
            
            # 3. Deduplication
            df, dedupe_report = dedupe.find_duplicates(df, ai_cleaner=ai)
            
            # 4. Job Roles
            df = job_role.map_job_roles(df)
            
            # 5. AI Cleaning (Smart Batch for Large Files)
            if len(df) > 1000:
                print(f"[Core] Large file detected ({len(df)} rows). Running AI on top 100 sample + rules for all.")
                # We run AI on a sample to populate "Preview" effectively without waiting hours
                # In a real heavy-duty app, this would be a background job.
                # Here we simulate the "withstand" capability by not choking.
                df_sample = df.head(100).copy()
                df_sample = ai.batch_clean(df_sample)
                df.update(df_sample)
            else:
                df = ai.batch_clean(df)

            # 6. Metrics
            metrics_data = metrics.calculate_metrics(df)
            
            # Save Results (Full Dataset)
            timestamp = int(time.time())
            output_filename = f"cleaned_{timestamp}"
            
            # Include dedupe report in analysis data
            analysis_data = {
                'metrics': metrics_data,
                'dedupe_report': dedupe_report,
                'data': df.head(200).to_dict(orient='records'),
                'total_rows': len(df),
                'timestamp': timestamp,
                'columns': df.columns.tolist()
            }
            
            output_csv = os.path.join(app.config['OUTPUT_FOLDER'], f"{output_filename}.csv")
            output_json = os.path.join(app.config['OUTPUT_FOLDER'], f"{output_filename}.json")
            output_qa_docx = os.path.join(app.config['OUTPUT_FOLDER'], f"{output_filename}_qa_report.docx")
            
            # Clean up internal columns before export
            export_df = df.drop(columns=['ai_suggestions'], errors='ignore')

            export_df.to_csv(output_csv, index=False)
            export_df.to_json(output_json, orient='records')
            
            # Generate QA Report Word Document
            doc = Document()
            
            # Title
            title = doc.add_heading('DATA QUALITY ASSURANCE REPORT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            metadata = doc.add_paragraph()
            metadata.add_run(f"Timestamp: {time.ctime(timestamp)}").italic = True
            metadata.add_run(f"\nTotal Records Processed: {len(df)}")
            
            # 1. Quality Scores
            doc.add_heading('QUALITY SCORES', level=1)
            
            # Calculations
            total_recs = len(df)
            col_count = metrics_data.get('column_count', 1)
            total_cells = total_recs * col_count
            invalid_count = metrics_data.get('total_field_issues', 0)
            error_rate = (invalid_count / total_cells * 100) if total_cells > 0 else 0
            
            score_p = doc.add_paragraph()
            score_p.add_run(f"Data Integrity Score: ").bold = True
            score_p.add_run(f"{(100 - error_rate):.2f}%")
            
            score_p.add_run(f"\nField Error Rate: ").bold = True
            score_p.add_run(f"{error_rate:.2f}%")
            
            score_p.add_run(f"\nOverall Quality Index: ").bold = True
            score_p.add_run(f"{metrics_data.get('after_score', 0)}% (Weighted)")
            
            # 2. Issue Breakdown
            doc.add_heading('DETAILED ISSUE BREAKDOWN', level=1)
            issue_summary = metrics_data.get('issue_summary', {})
            
            if issue_summary:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Issue Type'
                hdr_cells[1].text = 'Occurrences'
                hdr_cells[2].text = 'Impact (%)'
                
                # Bold headers
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                sorted_issues = sorted(issue_summary.items(), key=lambda x: x[1], reverse=True)
                for issue, count in sorted_issues:
                    row_cells = table.add_row().cells
                    row_cells[0].text = issue.replace('_', ' ').title()
                    row_cells[1].text = str(count)
                    pct = (count / total_recs) * 100
                    row_cells[2].text = f"{pct:.1f}%"
            else:
                doc.add_paragraph("No specific issues detected.")

            # 3. Field-Level Health
            doc.add_heading('FIELD-LEVEL HEALTH STATISTICS', level=1)
            column_health = metrics_data.get('column_health', {})
            if column_health:
                h_table = doc.add_table(rows=1, cols=4)
                h_table.style = 'Table Grid'
                h_hdr = h_table.rows[0].cells
                h_hdr[0].text = 'Field Name'
                h_hdr[1].text = 'Health (%)'
                h_hdr[2].text = 'Missing'
                h_hdr[3].text = 'Invalid'
                
                for cell in h_hdr:
                    for p in cell.paragraphs:
                        for r in p.runs: r.bold = True
                
                for col_name, stats in column_health.items():
                    r_cells = h_table.add_row().cells
                    r_cells[0].text = col_name.replace('_', ' ').title()
                    r_cells[1].text = f"{stats['health_pct']}%"
                    r_cells[2].text = str(stats['missing'])
                    r_cells[3].text = str(stats['invalid'])
            
            # 4. AI Correction Summary
            doc.add_heading('AI CORRECTION SUMMARY', level=1)
            remedy_summary = metrics_data.get('remedy_summary', {})
            if remedy_summary:
                doc.add_paragraph("The following automated corrections were applied to the dataset:")
                for remedy, count in remedy_summary.items():
                    doc.add_paragraph(f"• {remedy}: {count} fixes applied", style='List Bullet')
            else:
                doc.add_paragraph("No automated corrections were required for this dataset.")

            # 5. Data Issue Samples (Top 3 per issue)
            doc.add_heading('DATA ISSUE SAMPLES', level=1)
            if issue_summary:
                doc.add_paragraph("Below are representative samples of records containing identified issues:")
                for issue_key in sorted(issue_summary.keys()):
                    doc.add_heading(issue_key.replace('_', ' ').title(), level=2)
                    
                    # Exact match filtering for samples
                    # We use the original df for samples
                    sample_rows = df[df['issues'].fillna('').apply(lambda x: issue_key in [t.strip() for t in str(x).split(';') if t.strip()])].head(3)
                    
                    if not sample_rows.empty:
                        for idx, row in sample_rows.iterrows():
                            # Print a few key columns for context
                            p = doc.add_paragraph(style='List Bullet')
                            context_cols = [c for c in ['company_name', 'person_name', 'email', 'phone', 'website', 'country'] if c in df.columns]
                            context_parts = []
                            for c in context_cols:
                                if pd.notna(row[c]) and str(row[c]).strip():
                                    context_parts.append(f"{c.replace('_', ' ').title()}: {row[c]}")
                            p.add_run(" | ".join(context_parts) if context_parts else f"Row Index: {idx}")
                    else:
                        doc.add_paragraph("Sample not available (resolved in processing).")
            else:
                doc.add_paragraph("Dataset is healthy. No issues to sample.")
            
            doc.add_page_break()
            footer = doc.add_paragraph("\nReport generated by Data QA Copilot AI Engine.")
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.save(output_qa_docx)

            # Response: Send PREVIEW (Top 500 rows) + Full Count
            # This prevents browser crash on 100k rows
            # We explicitly slice the dataframe to 500 before converting to dict
            
            # Robust float cleaning for JSON
            # Convert to object to allow replacements
            if len(export_df) > 500:
                preview_df = export_df.head(500).copy()
            else:
                preview_df = export_df.copy()
            
            preview_df = preview_df.astype(object) 
            preview_df = preview_df.where(pd.notnull(preview_df), None)
            preview_rows = preview_df.to_dict(orient='records')
            
            # Prepare final response data
            response_data = {
                'success': True,
                'status': 'success',
                'metrics': metrics_data,
                'total_rows': len(df),
                'preview_limit': 500,
                'data': preview_rows,
                'timestamp': timestamp,
                'filename': output_filename,
                'qa_report': f"{output_filename}_qa_report.docx"
            }
            
            # recursive NaN cleaner for final safety
            import math
            import numpy as np
            def clean_for_json(obj):
                if isinstance(obj, (float, np.float64, np.float32)):
                    if math.isnan(obj) or math.isinf(obj):
                        return None
                    return float(obj)
                elif isinstance(obj, (int, np.integer)):
                    return int(obj)
                elif isinstance(obj, dict):
                    return {k: clean_for_json(v) for k, v in obj.items()}
                elif isinstance(obj, list):
                    return [clean_for_json(v) for v in obj]
                return obj
                
            safe_data = clean_for_json(response_data)
            
            return jsonify(safe_data)

        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"[ERROR] Processing failed: {e}")
            return jsonify({'error': f'Processing failed: {str(e)}'}), 500

@app.route('/api/download/<timestamp>/<format_type>', methods=['GET'])
def download_file(timestamp, format_type):
    # Security check on format
    if format_type not in ['csv', 'json', 'qa']:
        return "Invalid format", 400
        
    if format_type == 'qa':
        filename = f"cleaned_{timestamp}_qa_report.docx"
    else:
        filename = f"cleaned_{timestamp}.{format_type}"
        
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    
    if not os.path.exists(filepath):
        return "File not found", 404
        
    return send_file(filepath, as_attachment=True, download_name=filename)

@app.route('/api/chat', methods=['POST'])
def chat():
    data = request.json
    query = data.get('query', '')
    metrics_data = data.get('metrics', {})
    
    if not query:
        return jsonify({'error': 'No query provided'}), 400
        
    # Build a context string for the AI
    context = f"""
    - Overall Integrity Score: {100 - (metrics_data.get('total_field_issues',0) / (metrics_data.get('total_rows',1)*metrics_data.get('column_count',1)) * 100):.2f}%
    - Total Records: {metrics_data.get('total_rows', 0)}
    - Total Columns: {metrics_data.get('column_count', 0)}
    - Rows with Issues: {metrics_data.get('total_issues', 0)}
    - Specific Issues Found: {json.dumps(metrics_data.get('issue_summary', {}))}
    - Field-Level Health: {json.dumps(metrics_data.get('column_health', {}))}
    - AI Fixes Applied: {json.dumps(metrics_data.get('remedy_summary', {}))}
    """
    
    response = ai.chat(query, context)
    return jsonify({'response': response})

if __name__ == '__main__':
    app.run(debug=True, port=5000)
